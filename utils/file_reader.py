import docx
import openpyxl
import re
import logging
import os
from typing import List, Optional, Set
import threading
import time

logger = logging.getLogger(__name__)

# Thread-local storage for domain processing
_local = threading.local()
_local.domain_cache = {}

# Domain extraction patterns - compiled once for better performance
IP_PATTERN = re.compile(r'^(\d{1,3}\.){3}\d{1,3}$')
DOMAIN_PATTERN = re.compile(r'^([a-z0-9]([a-z0-9\-]{0,61}[a-z0-9])?\.)+[a-z]{2,}$')
LINE_SPLIT_PATTERN = re.compile(r'[\n\r]+')
PART_SPLIT_PATTERN = re.compile(r'[,\t ]+')
PROTOCOL_PATTERN = re.compile(r'^https?://')
URL_PATH_PATTERN = re.compile(r'/.*$')
TEXT_CLEANUP_PATTERN = re.compile(r'tekshirish natijalari:.*', flags=re.IGNORECASE)

# Cache for already cleaned domains
DOMAIN_CACHE_SIZE = 5000
domain_cache = {}


def clean_domain(domain: str) -> Optional[str]:
    """
    Domen nomini tozalash va validatsiya qilish.
    Subdomainlarni ham to'g'ri aniqlaydi (masalan: 'sur.ewe.test.uz').

    Args:
        domain: Tozalash kerak bo'lgan domen nomi

    Returns:
        Tozalangan domen yoki None (agar noto'g'ri formatda bo'lsa)
    """
    if not domain or not isinstance(domain, str):
        return None

    # Check cache first
    if domain in domain_cache:
        return domain_cache[domain]

    # Cleanup domain name
    domain = domain.strip().lower()

    # Remove text like "Tekshirish natijalari:" and similar
    domain = TEXT_CLEANUP_PATTERN.sub('', domain)

    # Remove protocols
    domain = PROTOCOL_PATTERN.sub('', domain)

    # Remove trailing slash and params
    domain = URL_PATH_PATTERN.sub('', domain)

    # One more whitespace check
    domain = domain.strip()

    if not domain:
        domain_cache[domain] = None
        return None

    # Check if it's an IP address
    if IP_PATTERN.match(domain):
        domain_cache[domain] = domain
        return domain

    # Check if it's a valid domain
    if DOMAIN_PATTERN.match(domain):
        domain_cache[domain] = domain
        return domain

    # Clean cache periodically
    if len(domain_cache) > DOMAIN_CACHE_SIZE:
        # Simple approach: just clear the whole cache when it gets too big
        domain_cache.clear()

    domain_cache[domain] = None
    return None


def extract_domains_from_text(text: str) -> List[str]:
    """
    Extract domains from text.

    Args:
        text: Text to extract domains from

    Returns:
        List of potential domains
    """
    if not text:
        return []

    domains = []

    # Split by newlines
    for line in LINE_SPLIT_PATTERN.split(text):
        line = line.strip()
        if not line:
            continue

        # Split by comma, tab or space
        parts = PART_SPLIT_PATTERN.split(line)
        domains.extend([part.strip() for part in parts if part.strip()])

    return domains


def read_docx_file(file_path: str, potential_domains: Set[str], max_domains: int) -> None:
    """Read domains from a .docx file"""
    try:
        doc = docx.Document(file_path)

        # Get from paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                domains_from_text = extract_domains_from_text(text)
                potential_domains.update(domains_from_text)
                if len(potential_domains) >= max_domains:
                    return

        # Get from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        domains_from_text = extract_domains_from_text(text)
                        potential_domains.update(domains_from_text)
                        if len(potential_domains) >= max_domains:
                            return
    except Exception as e:
        logger.error(f"Error reading docx file: {str(e)}")


def read_xlsx_file(file_path: str, potential_domains: Set[str], max_domains: int) -> None:
    """Read domains from a .xlsx file"""
    try:
        # Use read_only mode for better performance
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)

        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            # Process by rows for memory efficiency
            for row in sheet.rows:
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        text = cell.value.strip()
                        if text:
                            domains_from_text = extract_domains_from_text(text)
                            potential_domains.update(domains_from_text)
                            if len(potential_domains) >= max_domains:
                                return
    except Exception as e:
        logger.error(f"Error reading xlsx file: {str(e)}")


def read_file(file_path: str, max_domains: int = 5000) -> List[str]:
    """
    Read domains from file and format them correctly.
    Correctly identifies multi-level subdomains ('sur.ewe.test.uz').
    Improved to find all domains in file content.

    Args:
        file_path: Path to file
        max_domains: Maximum number of domains to process

    Returns:
        List of domains
    """
    start_time = time.time()
    potential_domains: Set[str] = set()
    processed_count = 0

    try:
        # Check if file exists
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return []

        # Read based on file type
        if file_path.endswith('.txt'):
            # Use a more efficient approach for text files
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    # Read in chunks for large files
                    chunk_size = 1024 * 1024  # 1MB chunks
                    text_buffer = ""

                    while True:
                        chunk = f.read(chunk_size)
                        if not chunk:
                            break

                        # Process any complete lines in this chunk
                        text_buffer += chunk
                        domains_from_text = extract_domains_from_text(text_buffer)
                        potential_domains.update(domains_from_text)

                        # Reset buffer to avoid memory build-up but keep any partial line
                        last_newline = text_buffer.rfind('\n')
                        if last_newline != -1:
                            text_buffer = text_buffer[last_newline + 1:]

                        if len(potential_domains) >= max_domains:
                            logger.warning(f"Reached maximum domains while reading file")
                            break

                # Process any remaining text in buffer
                if text_buffer:
                    domains_from_text = extract_domains_from_text(text_buffer)
                    potential_domains.update(domains_from_text)
            except Exception as e:
                logger.error(f"Error reading txt file: {str(e)}")

        elif file_path.endswith('.docx'):
            read_docx_file(file_path, potential_domains, max_domains)

        elif file_path.endswith('.xlsx'):
            read_xlsx_file(file_path, potential_domains, max_domains)

        # Clean and validate domains
        valid_domains = []
        for domain in potential_domains:
            cleaned = clean_domain(domain)
            if cleaned:
                valid_domains.append(cleaned)
                processed_count += 1

            # Check max domains
            if len(valid_domains) >= max_domains:
                logger.warning(f"Reached maximum domains ({max_domains}). Truncating list.")
                break

        # Remove duplicates and sort
        unique_domains = list(set(valid_domains))
        unique_domains.sort()

        end_time = time.time()
        logger.info(
            f"Read {len(potential_domains)} potential domains, {processed_count} processed, "
            f"{len(unique_domains)} unique in {end_time - start_time:.2f}s"
        )

        return unique_domains

    except Exception as e:
        logger.error(f"Error reading file: {str(e)}")
        return []